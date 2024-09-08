# Re-running the code to generate the docx file
from docx import Document

# Create a new Document
doc = Document()
doc.add_heading('Weekly Timetable for Deep Learning & Mathematics', 0)

# Add a note about daytime commitment and sleep priority
doc.add_paragraph(
    "Note: This timetable is designed to accommodate university and institute commitments during the daytime, ensuring ample time for sleep and practice in the evenings and weekends."
)

# Define the timetable
timetable = [
    ("Monday", [
        ("8:00 PM – 9:30 PM", "Mathematics (Linear Algebra)", "Focus on matrices, eigenvalues, eigenvectors, and matrix decomposition."),
        ("9:30 PM – 10:00 PM", "Break", "Take a short break."),
        ("10:00 PM – 11:30 PM", "Deep Learning (Neural Networks)", "Study basics of neural networks, activation functions, and backpropagation.")
    ]),
    ("Tuesday", [
        ("8:00 PM – 9:30 PM", "Mathematics (Calculus)", "Focus on derivatives, gradients, and integrals for optimization problems."),
        ("9:30 PM – 10:00 PM", "Break", "Break time."),
        ("10:00 PM – 11:30 PM", "Deep Learning (CNNs)", "Study convolutional neural networks (CNNs) and their applications.")
    ]),
    ("Wednesday", [
        ("8:00 PM – 9:30 PM", "Mathematics (Probability)", "Study probability distributions, random variables, and their applications in machine learning."),
        ("9:30 PM – 10:00 PM", "Break", "Break time."),
        ("10:00 PM – 11:30 PM", "Deep Learning (RNNs)", "Focus on recurrent neural networks (RNNs) and sequence modeling.")
    ]),
    ("Thursday", [
        ("8:00 PM – 9:30 PM", "Mathematics (Optimization)", "Learn gradient descent, SGD, and optimization methods."),
        ("9:30 PM – 10:00 PM", "Break", "Break time."),
        ("10:00 PM – 11:30 PM", "Deep Learning (Transfer Learning)", "Study transfer learning and fine-tuning models.")
    ]),
    ("Friday", [
        ("8:00 PM – 9:30 PM", "Mathematics (Statistics)", "Study statistics concepts like variance, hypothesis testing, and distributions."),
        ("9:30 PM – 10:00 PM", "Break", "Break time."),
        ("10:00 PM – 11:30 PM", "Deep Learning (Autoencoders)", "Study autoencoders and unsupervised learning techniques.")
    ]),
    ("Saturday", [
        ("8:00 PM – 11:00 PM", "Project Work", "Work on a larger deep learning project that incorporates math and deep learning concepts.")
    ]),
    ("Sunday", [
        ("9:00 AM – 12:00 PM", "Mathematics Review", "Review topics studied during the week and solve additional problems."),
        ("12:00 PM – 1:00 PM", "Break", "Take a break."),
        ("1:00 PM – 3:00 PM", "Project Review", "Review or complete any pending project work or topics.")
    ])
]

# Add the timetable to the document
for day, activities in timetable:
    doc.add_heading(day, level=1)
    for time, title, details in activities:
        doc.add_paragraph(f"{time}: {title}\nDetails: {details}", style='BodyText')

# Save the document
doc_path = "Deep_Learning_Mathematics_Timetable.docx"
doc.save(doc_path)

doc_path
